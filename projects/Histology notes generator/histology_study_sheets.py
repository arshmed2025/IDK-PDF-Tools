#!/usr/bin/env python3
"""
Histology Study Sheet Generator
================================
Reads a folder of labelled histology images and produces a Word document
where every page has the image, an editable heading from the filename,
and (optionally) bullets or ruled lines for key identification points.

USAGE
-----
1. One-time setup (in PowerShell or CMD):
       pip install python-docx Pillow

2. Drop this script into your images folder (alongside the .jpg/.png files).
   That's it — the default IMAGE_FOLDER = "" means "use my own folder".
   If you'd rather keep the script elsewhere, set IMAGE_FOLDER to the path
   of your images folder.

3. Optionally tweak the CONFIG block below:
   - BRANDING, fonts, font sizes, image quality
   - INCLUDE_NOTES, NUM_ITEMS, STYLE (bullets vs lines)
   - IMAGES_PER_PAGE (1, 2, 4, or 6)
   - ORDER_BY  (use the visual reorderer to generate order.txt)

4. Run by double-clicking the file, OR from terminal:
       python histology_study_sheets.py
"""

from pathlib import Path
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

# ============================================================
#  CONFIG  — edit these values
# ============================================================
# Folder containing your images. Leave empty ("") to use the folder this
# script is sitting in — handy if you put the script directly inside your
# images folder. Otherwise, give an absolute path:
#     IMAGE_FOLDER = r"C:\Users\YourName\Desktop\Histology Images"
IMAGE_FOLDER = ""

# Output filename. If just a name (no path), it's saved next to the script.
OUTPUT_FILE  = "Histology_Study_Sheets.docx"

BRANDING     = "IDK | Histology Notes"             # <-- your name / branding

# Image order in the document
ORDER_BY  = "custom"        # "name"          → A → Z (default)
                          # "name_reverse"  → Z → A
                          # "modified"      → newest file first
                          # "created"       → newest created first
                          # "custom"        → read order from ORDER_FILE
ORDER_FILE = "order.txt"  # used when ORDER_BY = "custom".  Place this
                          # plain-text file inside IMAGE_FOLDER, one
                          # filename per line in the order you want.
                          # Lines starting with "#" are ignored (comments).
                          # Any images not listed go at the end.

# Fonts — any font installed on your PC works (e.g. "Calibri", "Cambria",
# "Georgia", "Garamond", "Palatino Linotype", "Source Sans Pro", "Lato",
# "Merriweather", "Inter", "EB Garamond"). If the font isn't installed on
# whatever PC opens the doc, Word will substitute a fallback.
HEADING_FONT = "Montserrat"
BODY_FONT    = "Montserrat"

# Layout — how many images on each page
IMAGES_PER_PAGE = 1   # 1 = one big image per page (notes optional)
                      # 2 = two stacked vertically
                      # 4 = 2 x 2 grid
                      # 6 = 2 x 3 grid
                      # Note: INCLUDE_NOTES (bullets/lines) only works with 1.

# Notes area (only used when IMAGES_PER_PAGE = 1)
INCLUDE_NOTES = False         # True = label + bullets/lines below image
                              # False = just the image (no notes section)
STYLE       = "bullets"       # "bullets"  or  "lines"
NUM_ITEMS   = 5               # how many bullets (or lines) per page

# Font sizes (points). For multi-image layouts the heading auto-scales down.
HEADING_SIZE = 18             # filename heading at top of each image
LABEL_SIZE   = 16             # "Key identification points" label
BODY_SIZE    = 16             # size you type in (bullets / lines)
FOOTER_SIZE  = 12             # footer text

# Image sizing (inches) — used when IMAGES_PER_PAGE = 1.
# For multi-image layouts the dimensions are computed automatically.
MAX_IMG_WIDTH_INCHES  = 5.8
MAX_IMG_HEIGHT_INCHES = 4.6

# Image compression — HUGE impact on file size.
# Raw microscope photos are often 3-5 MB each; without this, a 91-image doc
# becomes ~300 MB. 200 DPI + quality 85 keeps images sharp in print and on
# screen while dropping total size by ~10x.
IMAGE_DPI     = 400           # target resolution (150-250 is the sweet spot)
IMAGE_QUALITY = 95            # JPEG quality 1-100 (85 ≈ visually lossless)

# Line style only
LINE_HEIGHT_CM = 0.75
# ============================================================


SUPPORTED_EXT = {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp", ".gif", ".webp"}

# Multi-image layouts: cols × rows, per-image max dims, heading-size scale.
LAYOUTS = {
    1: {"cols": 1, "rows": 1, "img_w": None, "img_h": None, "h_scale": 1.00},
    2: {"cols": 1, "rows": 2, "img_w": 6.4,  "img_h": 4.1,  "h_scale": 0.85},
    4: {"cols": 2, "rows": 2, "img_w": 3.1,  "img_h": 3.5,  "h_scale": 0.65},
    6: {"cols": 2, "rows": 3, "img_w": 3.1,  "img_h": 2.4,  "h_scale": 0.55},
}


# ------------------------------------------------------------
#  Low-level XML helpers
# ------------------------------------------------------------
def add_page_number_field(paragraph):
    """Insert a live PAGE field (auto-updating page number)."""
    run = paragraph.add_run()
    run.font.name = BODY_FONT
    run.font.size = Pt(FOOTER_SIZE)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def set_cell_bottom_border_only(cell, color="BFBFBF"):
    """Cell renders only a thin bottom border (writing line)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        borders.append(b)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    tc_pr.append(borders)


def remove_table_borders(table):
    """Strip all borders from a table so it's pure invisible layout."""
    tbl = table._element
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    # Remove any existing borders def
    for b in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(b)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "nil")
        borders.append(e)
    tbl_pr.append(borders)


def add_paragraph_top_border(paragraph, color="B0B0B0", size="6"):
    """Thin horizontal rule drawn above a paragraph (used in footer)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), size)
    top.set(qn("w:space"), "6")
    top.set(qn("w:color"), color)
    pBdr.append(top)
    pPr.append(pBdr)


def set_default_document_font(doc, font_name):
    """Make `font_name` the document's default so fresh typing inherits it."""
    styles_element = doc.styles.element
    rpr_default = styles_element.find(
        qn("w:docDefaults") + "/" + qn("w:rPrDefault") + "/" + qn("w:rPr")
    )
    if rpr_default is not None:
        for existing in rpr_default.findall(qn("w:rFonts")):
            rpr_default.remove(existing)
        rFonts = OxmlElement("w:rFonts")
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rFonts.set(qn(attr), font_name)
        rpr_default.insert(0, rFonts)

    try:
        normal = doc.styles["Normal"]
        normal.font.name = font_name
        rPr = normal.element.get_or_add_rPr()
        rFonts2 = rPr.find(qn("w:rFonts"))
        if rFonts2 is None:
            rFonts2 = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts2)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rFonts2.set(qn(attr), font_name)
    except KeyError:
        pass


def ensure_bullet_numbering(doc):
    """Create a bullet numbering definition once; return its numId."""
    try:
        numbering = doc.part.numbering_part.element
    except (NotImplementedError, AttributeError):
        tmp = doc.add_paragraph("", style="List Bullet")
        tmp._element.getparent().remove(tmp._element)
        numbering = doc.part.numbering_part.element

    abstract_id = "99"
    num_id      = "99"

    for tag, id_attr in (("w:abstractNum", "w:abstractNumId"),
                         ("w:num",         "w:numId")):
        for node in numbering.findall(qn(tag)):
            if node.get(qn(id_attr)) == abstract_id:
                numbering.remove(node)

    abs_num = OxmlElement("w:abstractNum")
    abs_num.set(qn("w:abstractNumId"), abstract_id)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    for tag, val in (("w:start", "1"),
                     ("w:numFmt", "bullet"),
                     ("w:lvlText", "\u2022"),
                     ("w:lvlJc", "left")):
        el = OxmlElement(tag); el.set(qn("w:val"), val); lvl.append(el)
    pPr_l = OxmlElement("w:pPr")
    ind   = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360"); ind.set(qn("w:hanging"), "360")
    pPr_l.append(ind); lvl.append(pPr_l)
    rPr_l = OxmlElement("w:rPr")
    rf    = OxmlElement("w:rFonts")
    rf.set(qn("w:ascii"), "Symbol"); rf.set(qn("w:hAnsi"), "Symbol")
    rf.set(qn("w:hint"), "default")
    rPr_l.append(rf); lvl.append(rPr_l)
    abs_num.append(lvl)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), num_id)
    aid = OxmlElement("w:abstractNumId"); aid.set(qn("w:val"), abstract_id)
    num.append(aid)

    first_num = numbering.find(qn("w:num"))
    if first_num is not None:
        first_num.addprevious(abs_num)
    else:
        numbering.append(abs_num)
    numbering.append(num)

    return int(num_id)


# ------------------------------------------------------------
#  Layout helpers
# ------------------------------------------------------------
def prepare_image(img_path, max_w_in=None, max_h_in=None):
    """Downsample + re-encode as JPEG so the embedded file stays small.

    Returns (BytesIO buffer, display_width_inches, display_height_inches).
    """
    if max_w_in is None: max_w_in = MAX_IMG_WIDTH_INCHES
    if max_h_in is None: max_h_in = MAX_IMG_HEIGHT_INCHES

    with Image.open(img_path) as im:
        orig_w, orig_h = im.size

        if im.mode in ("RGBA", "LA"):
            bg = Image.new("RGB", im.size, (255, 255, 255))
            bg.paste(im, mask=im.split()[-1])
            im = bg
        elif im.mode == "P":
            im = im.convert("RGBA")
            bg = Image.new("RGB", im.size, (255, 255, 255))
            bg.paste(im, mask=im.split()[-1])
            im = bg
        elif im.mode != "RGB":
            im = im.convert("RGB")

        aspect = orig_w / orig_h
        disp_w = max_w_in
        disp_h = disp_w / aspect
        if disp_h > max_h_in:
            disp_h = max_h_in
            disp_w = disp_h * aspect

        target_px_w = int(disp_w * IMAGE_DPI)
        target_px_h = int(disp_h * IMAGE_DPI)
        if im.width > target_px_w or im.height > target_px_h:
            im.thumbnail((target_px_w, target_px_h), Image.LANCZOS)

        buf = BytesIO()
        im.save(buf, format="JPEG",
                quality=IMAGE_QUALITY, optimize=True, progressive=True)
        buf.seek(0)

    return buf, Inches(disp_w), Inches(disp_h)


def prettify(name: str) -> str:
    """'thyroid_gland_01.jpg' -> 'Thyroid Gland 01'"""
    return Path(name).stem.replace("_", " ").replace("-", " ").strip().title()


def build_footer(section, branding_text):
    """Footer: thin divider line, italic branding (left), 'Page N' (right)."""
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = ""

    add_paragraph_top_border(p, color="C8C8C8", size="6")

    content_width_in = (section.page_width - section.left_margin - section.right_margin) / 914400
    pPr = p._p.get_or_add_pPr()
    for existing in pPr.findall(qn("w:tabs")):
        pPr.remove(existing)
    tabs = OxmlElement("w:tabs")
    for clear_pos in ("4680", "9360"):
        clr = OxmlElement("w:tab")
        clr.set(qn("w:val"), "clear"); clr.set(qn("w:pos"), clear_pos)
        tabs.append(clr)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right"); tab.set(qn("w:pos"), str(int(content_width_in * 1440)))
    tabs.append(tab)
    pPr.append(tabs)

    p.paragraph_format.space_before = Pt(6)

    left = p.add_run(branding_text + "\t")
    left.font.name  = BODY_FONT
    left.font.size  = Pt(FOOTER_SIZE)
    left.italic     = True
    left.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    pg_lbl = p.add_run("Page ")
    pg_lbl.font.name = BODY_FONT
    pg_lbl.font.size = Pt(FOOTER_SIZE)
    pg_lbl.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    add_page_number_field(p)


def add_bullet_paragraph(doc, num_id, text=""):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0"); numPr.append(ilvl)
    numId = OxmlElement("w:numId"); numId.set(qn("w:val"), str(num_id)); numPr.append(numId)
    pPr.append(numPr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(BODY_SIZE)
    return p


def add_image_page(doc, img_path, is_last, num_id):
    """One big image per page (with optional notes below)."""
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after  = Pt(8)
    r = h.add_run(prettify(img_path.name))
    r.bold = True
    r.font.size = Pt(HEADING_SIZE)
    r.font.name = HEADING_FONT
    r.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)

    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.paragraph_format.space_after = Pt(10)
    try:
        buf, w, h_ = prepare_image(img_path)
        pic_p.add_run().add_picture(buf, width=w, height=h_)
    except Exception as e:
        pic_p.add_run(f"[Could not embed image: {e}]")

    if INCLUDE_NOTES:
        lbl = doc.add_paragraph()
        lbl.paragraph_format.space_after = Pt(4)
        lrun = lbl.add_run("Key identification points")
        lrun.bold = True
        lrun.font.name = HEADING_FONT
        lrun.font.size = Pt(LABEL_SIZE)
        lrun.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

        if STYLE == "bullets":
            for _ in range(NUM_ITEMS):
                add_bullet_paragraph(doc, num_id)
        else:
            table = doc.add_table(rows=NUM_ITEMS, cols=1)
            for row in table.rows:
                row.height = Cm(LINE_HEIGHT_CM)
                cell = row.cells[0]
                set_cell_bottom_border_only(cell)
                para = cell.paragraphs[0]
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after  = Pt(0)

    if not is_last:
        doc.add_page_break()


def add_image_grid_page(doc, img_paths, is_last):
    """Lay out 2/4/6 images on one page using an invisible table."""
    layout = LAYOUTS[IMAGES_PER_PAGE]
    cols, rows = layout["cols"], layout["rows"]
    img_w, img_h = layout["img_w"], layout["img_h"]
    h_size = max(9, round(HEADING_SIZE * layout["h_scale"]))

    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for idx in range(rows * cols):
        r, c = divmod(idx, cols)
        cell = table.rows[r].cells[c]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        # Reset the default empty paragraph in the cell
        first_para = cell.paragraphs[0]
        first_para.text = ""
        first_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        first_para.paragraph_format.space_before = Pt(0)
        first_para.paragraph_format.space_after  = Pt(2)

        if idx >= len(img_paths):
            continue   # leave trailing cells empty

        img_path = img_paths[idx]

        # Heading run inside the existing first paragraph
        run = first_para.add_run(prettify(img_path.name))
        run.bold = True
        run.font.size = Pt(h_size)
        run.font.name = HEADING_FONT
        run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)

        # Image paragraph
        pic_p = cell.add_paragraph()
        pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_p.paragraph_format.space_before = Pt(0)
        pic_p.paragraph_format.space_after  = Pt(6)
        try:
            buf, w, h_ = prepare_image(img_path, img_w, img_h)
            pic_p.add_run().add_picture(buf, width=w, height=h_)
        except Exception as e:
            pic_p.add_run(f"[Could not embed: {e}]")

    remove_table_borders(table)

    if not is_last:
        doc.add_page_break()


def order_images(images, folder):
    """Return `images` reordered according to ORDER_BY."""
    if ORDER_BY == "name":
        return sorted(images, key=lambda p: p.name.lower())
    if ORDER_BY == "name_reverse":
        return sorted(images, key=lambda p: p.name.lower(), reverse=True)
    if ORDER_BY == "modified":
        return sorted(images, key=lambda p: p.stat().st_mtime, reverse=True)
    if ORDER_BY == "created":
        return sorted(images, key=lambda p: p.stat().st_ctime, reverse=True)
    if ORDER_BY == "custom":
        order_path = folder / ORDER_FILE

        if not order_path.is_file():
            with open(order_path, "w", encoding="utf-8") as f:
                f.write("# Edit this file to set the page order.\n")
                f.write("# One filename per line. Lines starting with '#' are ignored.\n")
                f.write("# You can group with comments like:  # ----- THYROID -----\n\n")
                for p in sorted(images, key=lambda p: p.name.lower()):
                    f.write(p.name + "\n")
            raise SystemExit(
                f"\n[STARTER FILE CREATED]\n"
                f"  {order_path}\n"
                f"Open it, drag lines into the order you want, save, "
                f"then run this script again."
            )

        listed = []
        with open(order_path, encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith("#"):
                    listed.append(line)

        by_name = {p.name.lower(): p for p in images}
        ordered, used = [], set()
        for name in listed:
            key = name.lower()
            if key in by_name and key not in used:
                ordered.append(by_name[key])
                used.add(key)
            else:
                matches = [p for p in images
                           if p.stem.lower() == key and p.name.lower() not in used]
                if matches:
                    ordered.append(matches[0])
                    used.add(matches[0].name.lower())
                else:
                    print(f"  [WARN] '{name}' listed in {ORDER_FILE} but not found in folder")

        leftover = sorted(
            (p for p in images if p.name.lower() not in used),
            key=lambda p: p.name.lower()
        )
        if leftover:
            print(f"  [INFO] {len(leftover)} unlisted image(s) appended at end")
        return ordered + leftover

    raise SystemExit(f"[ERROR] Unknown ORDER_BY value: {ORDER_BY!r}")


def main():
    if IMAGES_PER_PAGE not in LAYOUTS:
        raise SystemExit(f"[ERROR] IMAGES_PER_PAGE must be one of {sorted(LAYOUTS)}; "
                         f"got {IMAGES_PER_PAGE}")
    notes_active = INCLUDE_NOTES and IMAGES_PER_PAGE == 1
    if INCLUDE_NOTES and IMAGES_PER_PAGE != 1:
        print(f"[INFO] INCLUDE_NOTES is on but IMAGES_PER_PAGE = {IMAGES_PER_PAGE}; "
              f"notes are skipped (only used with 1 image per page).")

    script_dir = Path(__file__).parent.resolve()

    if IMAGE_FOLDER.strip():
        folder = Path(IMAGE_FOLDER).expanduser()
        if not folder.is_absolute():
            folder = (script_dir / folder).resolve()
    else:
        folder = script_dir
    print(f"Looking for images in: {folder}")

    if not folder.is_dir():
        raise SystemExit(f"[ERROR] Folder not found: {folder}\n"
                         f"Edit IMAGE_FOLDER at the top of the script, "
                         f"or place this script inside your image folder.")

    images = [p for p in folder.iterdir() if p.suffix.lower() in SUPPORTED_EXT]
    if not images:
        raise SystemExit(f"[ERROR] No supported images found in: {folder}")

    images = order_images(images, folder)
    n_pages = -(-len(images) // IMAGES_PER_PAGE)   # ceil division
    print(f"Found {len(images)} image(s). Order: {ORDER_BY}. "
          f"Layout: {IMAGES_PER_PAGE}/page → {n_pages} page(s). Building...")

    doc = Document()
    set_default_document_font(doc, BODY_FONT)

    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)
        build_footer(section, BRANDING)

    num_id = ensure_bullet_numbering(doc) if (notes_active and STYLE == "bullets") else None

    if IMAGES_PER_PAGE == 1:
        for i, img in enumerate(images):
            print(f"  [{i+1}/{len(images)}] {img.name}")
            add_image_page(doc, img, is_last=(i == len(images) - 1), num_id=num_id)
    else:
        n = IMAGES_PER_PAGE
        for page_idx in range(n_pages):
            chunk = images[page_idx * n : (page_idx + 1) * n]
            names = ", ".join(img.name for img in chunk)
            print(f"  [page {page_idx+1}/{n_pages}] {names}")
            add_image_grid_page(doc, chunk, is_last=(page_idx == n_pages - 1))

    out_path = Path(OUTPUT_FILE).expanduser()
    if not out_path.is_absolute():
        out_path = script_dir / out_path
    out_path = out_path.resolve()
    doc.save(out_path)
    print(f"\nDone. Saved: {out_path}")
    if IMAGES_PER_PAGE == 1:
        notes_desc = f"{STYLE} ({NUM_ITEMS})" if notes_active else "image only"
        print(f"Pages: {n_pages}  |  1 img/page  |  Notes: {notes_desc}  |  Font: {BODY_FONT}")
    else:
        print(f"Pages: {n_pages}  |  {IMAGES_PER_PAGE} imgs/page  |  Font: {BODY_FONT}")


if __name__ == "__main__":
    import sys
    try:
        main()
    except KeyboardInterrupt:
        print("\nCancelled.")
    except SystemExit as e:
        if isinstance(e.code, str):
            print(e.code, file=sys.stderr)
    except Exception:
        import traceback
        traceback.print_exc()
    try:
        input("\nPress Enter to close...")
    except EOFError:
        pass
